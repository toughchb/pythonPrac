from docx import Document

document = Document()

document.add_heading('Document Title',1)

document.save('demo.docx')
#doc = docx.Document('')
