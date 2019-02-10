from docx import Document
from googletrans import Translator
translator = Translator()

doc = Document('input3.docx')
for i in doc.paragraphs:
    output = translator.translate(i.text,src='en', dest='ko')
    doc.add_paragraph(output.text)
    #print(output.text)

doc.save('output2.docx')
