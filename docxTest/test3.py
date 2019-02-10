from docx import Document
from googletrans import Translator
import sys

translator = Translator()

doc = Document('input2.docx')
doc.paragraphs
#print("#1 :" + str(doc.paragraphs))
print(doc.paragraphs[0])
bb = ""
for n in doc.paragraphs:
    bb += n.text
    if bb[-2:] != '.\n':
        bb.rstrip('\n')
print(bb)
output = []
output.append(bb)
#for i in bb:
#if bb[-2:] != '.\n':
#    output = bb.rstrip('\n')

    #output = translator.translate(i.text,src='en', dest='ko')
#    doc.add_paragraph(output.text)
print(output)
#
#doc.save('output2.docx')
